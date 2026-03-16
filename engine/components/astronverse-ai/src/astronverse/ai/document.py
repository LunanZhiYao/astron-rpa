"""Document AI utilities for expansion and summarization prompts."""

# encoding: UTF-8

import copy
import uuid
from pathlib import Path
from astronverse.actionlib import AtomicFormType, AtomicFormTypeMeta, DynamicsItem
from astronverse.actionlib.atomic import atomicMg
from astronverse.actionlib.logger import logger
from astronverse.ai import InputType, LLMModelTypes, OutputType, SaveType, TemplateType
from astronverse.ai.api.llm import chat_normal
from astronverse.ai.chat import ChatAI
from astronverse.ai.prompt.document import (
    PROMPT_SENTENCE_EXTEND,
    PROMPT_SENTENCE_REDUCE,
    PROMPT_THEME_EXTEND,
    PROMPT_LAYOUT,
    PROMPT_AUTOARRANGE_NOTES_TEXT,
    PROMPT_AUTOARRANGE_NOTES_MINDMAP,
    PROMPT_DOCUMENT_SUMMARY,
    PROMPT_DOCUMENT_ANALYSIS,
    PROMPT_DOCUMENT_POLISHING,
)
from astronverse.ai.utils.str import replace_keyword


class DocumentAI:
    """Provide document-oriented AI utilities: theme expansion, sentence expansion and reduction."""

    @staticmethod
    @atomicMg.atomic(
        "DocumentAI",
        inputList=[
            atomicMg.param(
                "prompt",
                formType=AtomicFormTypeMeta(
                    type="INPUT_PYTHON_TEXTAREAMODAL_VARIABLE",
                ),
            ),
            atomicMg.param(
                "theme",
                formType=AtomicFormTypeMeta(
                    type="INPUT_VARIABLE_PYTHON",
                ),
            ),
        ],
        outputList=[atomicMg.param("theme_expanded_res", types="Str")],
    )
    def theme_expand(
        model_select: LLMModelTypes = LLMModelTypes.DEEPSEEK_V3_2,
        prompt: str = PROMPT_THEME_EXTEND,
        theme: str = "",
    ) -> str:
        """
        主题扩写
        Args:
            - model_select(LLMModelTypes): 模型选择
            - prompt(str): 提示词
            - theme(str): 主题
        Return:
            `str`, 扩写后的文本
        """
        model = model_select.value if isinstance(model_select, LLMModelTypes) else model_select

        user_input = f"{prompt}\n\n主题：{theme}"

        return chat_normal(user_input=user_input, system_input="", model=model)

    @staticmethod
    @atomicMg.atomic(
        "DocumentAI",
        inputList=[
            atomicMg.param(
                "document_path",
                dynamics=[
                    DynamicsItem(
                        key="$this.document_path.show",
                        expression="return $this.input_mode.value == '{}'".format(InputType.FILE.value),
                    )
                ],
                formType=AtomicFormTypeMeta(
                    type=AtomicFormType.INPUT_VARIABLE_PYTHON_FILE.value,
                    params={"filters": [".doc", ".docx", ".ppt", ".pptx", ".pdf"], "file_type": "file"},
                ),
            ),
            atomicMg.param(
                "note_content",
                dynamics=[
                    DynamicsItem(
                        key="$this.note_content.show",
                        expression="return $this.input_mode.value == '{}'".format(InputType.TEXT.value),
                    )
                ],
                formType=AtomicFormTypeMeta(
                    type="INPUT_PYTHON_TEXTAREAMODAL_VARIABLE",
                ),
            ),
            atomicMg.param(
                "image_output_path",
                dynamics=[
                    DynamicsItem(
                        key="$this.image_output_path.show",
                        expression="return $this.output_type.value == '{}'".format(OutputType.MINDMAP.value),
                    )
                ],
                formType=AtomicFormTypeMeta(
                    type="INPUT_VARIABLE_PYTHON_FILE",
                    params={"file_type": "folder"},
                ),
            ),
            atomicMg.param(
                "prompt_text",
                dynamics=[
                    DynamicsItem(
                        key="$this.prompt_text.show",
                        expression="return $this.output_type.value == '{}'".format(OutputType.TEXT.value),
                    )
                ],
                formType=AtomicFormTypeMeta(
                    type="INPUT_PYTHON_TEXTAREAMODAL_VARIABLE",
                ),
            ),
            atomicMg.param(
                "prompt_mindmap",
                dynamics=[
                    DynamicsItem(
                        key="$this.prompt_mindmap.show",
                        expression="return $this.output_type.value == '{}'".format(OutputType.MINDMAP.value),
                    )
                ],
                formType=AtomicFormTypeMeta(
                    type="INPUT_PYTHON_TEXTAREAMODAL_VARIABLE",
                ),
            ),
        ],
        outputList=[atomicMg.param("AI_text_organization_res", types="Str")],
    )
    def organize_notes(
        model_select: LLMModelTypes = LLMModelTypes.DEEPSEEK_V3_2,
        prompt_text: str = PROMPT_AUTOARRANGE_NOTES_TEXT,
        prompt_mindmap: str = PROMPT_AUTOARRANGE_NOTES_MINDMAP,
        input_mode: InputType = InputType.TEXT,
        document_path: str = "",
        note_content: str = "",
        output_type: OutputType = OutputType.TEXT,
        image_output_path: str = "",
    ) -> str:
        """
        自动整理笔记
        Args:
            - model_select(LLMModelTypes): 模型选择
            - prompt_text(str): 文本输出提示词
            - prompt_mindmap(str): 思维导图输出提示词
            - input_mode(InputType): 输入方式（file/text）
            - document_path(str): 文档路径
            - note_content(str): 笔记内容
            - output_type(OutputType): 输出类型（text/mindmap）
            - image_output_path(str): 图片输出路径
        Return:
            `str`, 整理后的笔记或图片路径
        """
        import re
        from pathlib import Path

        import requests
        from astronverse.ai.utils.extract import FileExtractor

        model = model_select.value if isinstance(model_select, LLMModelTypes) else model_select

        # 处理 input_mode 枚举
        input_mode_value = input_mode.value if isinstance(input_mode, InputType) else input_mode

        # 获取笔记内容
        if input_mode_value == InputType.FILE.value:
            content = FileExtractor(document_path).extract_text()
        else:
            content = note_content

        # 处理 output_type 枚举
        output_type_value = output_type.value if isinstance(output_type, OutputType) else output_type

        # 根据输出类型选择提示词
        if output_type_value == OutputType.MINDMAP.value:
            user_input = f"{prompt_mindmap}\n\n笔记内容：\n{content}"
        else:
            user_input = f"{prompt_text}\n\n笔记内容：\n{content}"

        # 调用 LLM
        result = chat_normal(user_input=user_input, system_input="", model=model)

        # 处理输出
        if output_type_value == OutputType.MINDMAP.value:
            # 提取 Mermaid 代码
            mermaid_match = re.search(r"```mermaid\s*(.*?)\s*```", result, re.DOTALL)
            if mermaid_match:
                mermaid_code = mermaid_match.group(1).strip()
            else:
                # 如果没有代码块，尝试直接使用整个结果
                mermaid_code = result.strip()

            # 使用 Kroki API 渲染为图片（更稳定的替代方案）
            kroki_url = "https://kroki.io/mermaid/png"

            response = requests.post(
                kroki_url, json={"diagram_source": mermaid_code}, headers={"Content-Type": "application/json"}
            )
            if response.status_code == 200:
                # 将 image_output_path 视为目录
                save_dir = Path(image_output_path)
                # 生成随机ID（例如取UUID前8位）
                random_id = str(uuid.uuid4())[:8]
                # 构造文件名：mindmap_随机ID.png
                filename = f"mindmap_{random_id}.png"
                save_path = save_dir / filename

                # 确保目录存在
                save_dir.mkdir(parents=True, exist_ok=True)
                # 保存图片
                save_path.write_bytes(response.content)

                return f"思维导图已保存至：{save_path}"
            else:
                return f"思维导图渲染失败（状态码：{response.status_code}），Mermaid代码：\n{mermaid_code}"
        else:
            return result

    @staticmethod
    @atomicMg.atomic(
        "DocumentAI",
        inputList=[
            atomicMg.param(
                "document_path",
                dynamics=[
                    DynamicsItem(
                        key="$this.document_path.show",
                        expression="return $this.input_mode.value == '{}'".format(InputType.FILE.value),
                    )
                ],
                formType=AtomicFormTypeMeta(
                    type=AtomicFormType.INPUT_VARIABLE_PYTHON_FILE.value,
                    params={"filters": [".doc", ".docx", ".ppt", ".pptx", ".pdf"], "file_type": "file"},
                ),
            ),
            atomicMg.param(
                "document_content",
                dynamics=[
                    DynamicsItem(
                        key="$this.document_content.show",
                        expression="return $this.input_mode.value == '{}'".format(InputType.TEXT.value),
                    )
                ],
                formType=AtomicFormTypeMeta(
                    type="INPUT_PYTHON_TEXTAREAMODAL_VARIABLE",
                ),
            ),
            atomicMg.param(
                "prompt",
                formType=AtomicFormTypeMeta(
                    type="INPUT_PYTHON_TEXTAREAMODAL_VARIABLE",
                ),
            ),
        ],
        outputList=[atomicMg.param("document_summary_res", types="Str")],
    )
    def summarize_document(
        model_select: LLMModelTypes = LLMModelTypes.DEEPSEEK_V3_2,
        prompt: str = PROMPT_DOCUMENT_SUMMARY,
        input_mode: InputType = InputType.TEXT,
        document_path: str = "",
        document_content: str = "",
    ) -> str:
        """
        文档总结
        Args:
            - model_select(LLMModelTypes): 模型选择
            - prompt(str): 提示词
            - input_mode(InputType): 输入方式（file/text）
            - document_path(str): 文档路径
            - document_content(str): 文档内容
        Return:
            `str`, 总结后的文本
        """
        from astronverse.ai.utils.extract import FileExtractor

        model = model_select.value if isinstance(model_select, LLMModelTypes) else model_select

        # 处理 input_mode 枚举
        input_mode_value = input_mode.value if isinstance(input_mode, InputType) else input_mode

        # 获取文档内容
        if input_mode_value == InputType.FILE.value:
            content = FileExtractor(document_path).extract_text()
        else:
            content = document_content

        user_input = f"{prompt}\n\n文档内容：\n{content}"

        return chat_normal(user_input=user_input, system_input="", model=model)

    @staticmethod
    @atomicMg.atomic(
        "DocumentAI",
        inputList=[
            atomicMg.param(
                "document_path",
                dynamics=[
                    DynamicsItem(
                        key="$this.document_path.show",
                        expression="return $this.input_mode.value == '{}'".format(InputType.FILE.value),
                    )
                ],
                formType=AtomicFormTypeMeta(
                    type=AtomicFormType.INPUT_VARIABLE_PYTHON_FILE.value,
                    params={"filters": [".doc", ".docx", ".ppt", ".pptx", ".pdf"], "file_type": "file"},
                ),
            ),
            atomicMg.param(
                "document_content",
                dynamics=[
                    DynamicsItem(
                        key="$this.document_content.show",
                        expression="return $this.input_mode.value == '{}'".format(InputType.TEXT.value),
                    )
                ],
                formType=AtomicFormTypeMeta(
                    type="INPUT_PYTHON_TEXTAREAMODAL_VARIABLE",
                ),
            ),
            atomicMg.param(
                "prompt",
                formType=AtomicFormTypeMeta(
                    type="INPUT_PYTHON_TEXTAREAMODAL_VARIABLE",
                ),
            ),
        ],
        outputList=[atomicMg.param("document_analysis_res", types="Str")],
    )
    def analyze_document(
        model_select: LLMModelTypes = LLMModelTypes.DEEPSEEK_V3_2,
        prompt: str = PROMPT_DOCUMENT_ANALYSIS,
        input_mode: InputType = InputType.TEXT,
        document_path: str = "",
        document_content: str = "",
    ) -> str:
        """
        文档分析
        Args:
            - model_select(LLMModelTypes): 模型选择
            - prompt(str): 提示词
            - input_mode(InputType): 输入方式（file/text）
            - document_path(str): 文档路径
            - document_content(str): 文档内容
        Return:
            `str`, 分析后的文本
        """
        from astronverse.ai.utils.extract import FileExtractor

        model = model_select.value if isinstance(model_select, LLMModelTypes) else model_select

        # 处理 input_mode 枚举
        input_mode_value = input_mode.value if isinstance(input_mode, InputType) else input_mode

        # 获取文档内容
        if input_mode_value == InputType.FILE.value:
            content = FileExtractor(document_path).extract_text()
        else:
            content = document_content

        user_input = f"{prompt}\n\n文档内容：\n{content}"

        return chat_normal(user_input=user_input, system_input="", model=model)

    @staticmethod
    @atomicMg.atomic(
        "DocumentAI",
        inputList=[
            atomicMg.param(
                "document_path",
                dynamics=[
                    DynamicsItem(
                        key="$this.document_path.show",
                        expression="return $this.input_mode.value == '{}'".format(InputType.FILE.value),
                    )
                ],
                formType=AtomicFormTypeMeta(
                    type=AtomicFormType.INPUT_VARIABLE_PYTHON_FILE.value,
                    params={"filters": [".doc", ".docx", ".ppt", ".pptx", ".pdf"], "file_type": "file"},
                ),
            ),
            atomicMg.param(
                "document_content",
                dynamics=[
                    DynamicsItem(
                        key="$this.document_content.show",
                        expression="return $this.input_mode.value == '{}'".format(InputType.TEXT.value),
                    )
                ],
                formType=AtomicFormTypeMeta(
                    type="INPUT_PYTHON_TEXTAREAMODAL_VARIABLE",
                ),
            ),
            atomicMg.param(
                "prompt",
                formType=AtomicFormTypeMeta(
                    type="INPUT_PYTHON_TEXTAREAMODAL_VARIABLE",
                ),
            ),
        ],
        outputList=[atomicMg.param("AI_text_polishing_res", types="Str")],
    )
    def polish_document(
        model_select: LLMModelTypes = LLMModelTypes.DEEPSEEK_V3_2,
        prompt: str = PROMPT_DOCUMENT_POLISHING,
        input_mode: InputType = InputType.TEXT,
        document_path: str = "",
        document_content: str = "",
    ) -> str:
        """
        文档润色
        Args:
            - model_select(LLMModelTypes): 模型选择
            - prompt(str): 提示词
            - input_mode(InputType): 输入方式（file/text）
            - document_path(str): 文档路径
            - document_content(str): 文档内容
        Return:
            `str`, 润色后的文本
        """
        from astronverse.ai.utils.extract import FileExtractor

        model = model_select.value if isinstance(model_select, LLMModelTypes) else model_select

        # 处理 input_mode 枚举
        input_mode_value = input_mode.value if isinstance(input_mode, InputType) else input_mode

        # 获取文档内容
        if input_mode_value == InputType.FILE.value:
            content = FileExtractor(document_path).extract_text()
        else:
            content = document_content

        user_input = f"{prompt}\n\n文档内容：\n{content}"

        return chat_normal(user_input=user_input, system_input="", model=model)

    @staticmethod
    @atomicMg.atomic(
        "DocumentAI",
        inputList=[
            atomicMg.param(
                "document_path",
                dynamics=[
                    DynamicsItem(
                        key="$this.document_path.show",
                        expression="return $this.input_mode.value == '{}'".format(InputType.FILE.value),
                    )
                ],
                formType=AtomicFormTypeMeta(
                    type=AtomicFormType.INPUT_VARIABLE_PYTHON_FILE.value,
                    params={"filters": [".doc", ".docx", ".txt", ".json"], "file_type": "file"},
                ),
            ),
            atomicMg.param(
                "document_content",
                dynamics=[
                    DynamicsItem(
                        key="$this.document_content.show",
                        expression="return $this.input_mode.value == '{}'".format(InputType.TEXT.value),
                    )
                ],
                formType=AtomicFormTypeMeta(
                    type="INPUT_PYTHON_TEXTAREAMODAL_VARIABLE",
                ),
            ),
            atomicMg.param(
                "template_select",
                formType=AtomicFormTypeMeta(type=AtomicFormType.SELECT.value),
            ),
            atomicMg.param(
                "save_type",
                formType=AtomicFormTypeMeta(type=AtomicFormType.RADIO.value),
            ),
            atomicMg.param(
                "output_path",
                dynamics=[
                    DynamicsItem(
                        key="$this.output_path.show",
                        expression="return $this.save_type.value == '{}'".format(SaveType.SAVE_AS.value),
                    )
                ],
                formType=AtomicFormTypeMeta(
                    type="INPUT_VARIABLE_PYTHON_FILE",
                    params={"file_type": "folder"},
                ),
            ),
            atomicMg.param(
                "output_filename",
                dynamics=[
                    DynamicsItem(
                        key="$this.output_filename.show",
                        expression="return $this.save_type.value == '{}'".format(SaveType.SAVE_AS.value),
                    )
                ],
                formType=AtomicFormTypeMeta(
                    type="INPUT_VARIABLE_PYTHON",
                ),
            ),
        ],
        outputList=[atomicMg.param("AI_text_formatting_res", types="Str")],
    )
    def format_document(
        input_mode: InputType = InputType.TEXT,
        document_path: str = "",
        document_content: str = "",
        template_select: TemplateType = TemplateType.GENERAL,
        save_type: SaveType = SaveType.SAVE,
        output_path: str = "",
        output_filename: str = "",
    ) -> str:
        """
        AI排版
        Args:
            - input_mode(InputType): 输入方式（file/text）
            - document_path(str): 文档路径
            - document_content(str): 文档内容
            - template_select(TemplateType): 模板选择（通用公文/学术论文/合同协议）
            - save_type(SaveType): 保存类型（保存/另存为）
            - output_path(str): 文档输出路径
            - output_filename(str): 文档输出文件名
        Return:
            `str`, 格式化后的文档路径
        """
        import json
        from pathlib import Path

        from astronverse.ai.utils.extract import FileExtractor
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt, RGBColor

        model_select: LLMModelTypes = LLMModelTypes.DEEPSEEK_V3_2
        model = model_select.value if isinstance(model_select, LLMModelTypes) else model_select

        # 处理 input_mode 枚举
        input_mode_value = input_mode.value if isinstance(input_mode, InputType) else input_mode

        # 处理 save_type 枚举
        save_type_value = save_type.value if isinstance(save_type, SaveType) else save_type

        # 处理 template_select 枚举
        template_select_value = template_select.value if isinstance(template_select, TemplateType) else template_select

        # 获取文档内容
        if input_mode_value == InputType.FILE.value:
            content = FileExtractor(document_path).extract_text()
        else:
            content = document_content

        # 构建结构化标注提示词
        structure_prompt = (
            PROMPT_LAYOUT
            + f"""
                            文档内容：
                            {content}
                            """
        )

        # 调用LLM进行结构化标注
        structure_result = chat_normal(user_input=structure_prompt, system_input="", model=model)

        # 解析JSON结果
        try:
            # 尝试提取JSON内容
            import re

            json_match = re.search(r"\{.*\}", structure_result, re.DOTALL)
            if json_match:
                structure_data = json.loads(json_match.group(0))
            else:
                structure_data = json.loads(structure_result)
        except json.JSONDecodeError:
            # 如果解析失败，使用简单的段落分割
            structure_data = {
                "elements": [
                    {"order_id": i, "type": "paragraph", "content": para, "level": 0}
                    for i, para in enumerate(content.split("\n\n"))
                    if para.strip()
                ]
            }

        logger.info(f"structure_data: {structure_data}")

        # 创建Word文档
        doc = Document()

        # 应用模板样式
        if template_select_value == TemplateType.GENERAL.value:
            # 设置页边距
            sections = doc.sections
            for section in sections:
                section.page_height = Cm(29.7)  # A4高度
                section.page_width = Cm(21)  # A4宽度
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(2.54)
                section.right_margin = Cm(2.54)

            # 应用元素样式
            for element in structure_data.get("elements", []):
                elem_type = element.get("type", "paragraph")
                elem_content = element.get("content", "")
                elem_level = element.get("level", 0)

                if elem_type == "title":
                    p = doc.add_heading(elem_content, level=0)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif elem_type.startswith("heading"):
                    level = int(elem_type[-1]) if elem_type[-1].isdigit() else elem_level
                    p = doc.add_heading(elem_content, level=level)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif elem_type == "list_item":
                    p = doc.add_paragraph(elem_content, style="List Bullet")
                else:
                    p = doc.add_paragraph(elem_content)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    # 设置首行缩进
                    p.paragraph_format.first_line_indent = Cm(0.74)
                    p.paragraph_format.line_spacing = 1.5

        elif template_select_value == TemplateType.ACADEMIC.value:
            # 学术论文样式
            sections = doc.sections
            for section in sections:
                section.page_height = Cm(29.7)
                section.page_width = Cm(21)
                section.top_margin = Cm(2.5)
                section.bottom_margin = Cm(2.5)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)

            for element in structure_data.get("elements", []):
                elem_type = element.get("type", "paragraph")
                elem_content = element.get("content", "")

                if elem_type == "title":
                    p = doc.add_heading(elem_content, level=0)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif elem_type.startswith("heading"):
                    level = int(elem_type[-1]) if elem_type[-1].isdigit() else 1
                    p = doc.add_heading(elem_content, level=level)
                else:
                    p = doc.add_paragraph(elem_content)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.first_line_indent = Cm(0.74)
                    p.paragraph_format.line_spacing = 1.5

        elif template_select_value == TemplateType.CONTRACT.value:
            # 合同协议样式
            sections = doc.sections
            for section in sections:
                section.page_height = Cm(29.7)
                section.page_width = Cm(21)
                section.top_margin = Cm(2.5)
                section.bottom_margin = Cm(2.5)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)

            for element in structure_data.get("elements", []):
                elem_type = element.get("type", "paragraph")
                elem_content = element.get("content", "")

                if elem_type == "title":
                    p = doc.add_heading(elem_content, level=0)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif elem_type.startswith("heading"):
                    p = doc.add_paragraph(elem_content)
                    p.runs[0].bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p = doc.add_paragraph(elem_content)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.line_spacing = 1.5

        # 保存文档
        if save_type_value == SaveType.SAVE.value and input_mode_value == InputType.FILE.value:
            # 覆盖原文件
            output_file = document_path
        else:
            # 另存为
            if not output_filename:
                # 生成随机ID（例如取UUID前8位）
                random_id = str(uuid.uuid4())[:8]
                # 构造文件名：mindmap_随机ID.png
                output_filename = f"{random_id}"
            if not output_filename.endswith(".docx"):
                output_filename += ".docx"
            output_file = str(Path(output_path) / output_filename)

        doc.save(output_file)

        return f"文档已成功格式化并保存至：{output_file}"

    @staticmethod
    @atomicMg.atomic(
        "DocumentAI",
        inputList=[
            atomicMg.param(
                "prompt",
                formType=AtomicFormTypeMeta(
                    type="INPUT_PYTHON_TEXTAREAMODAL_VARIABLE",
                ),
            ),
            atomicMg.param(
                "paragraph_text",
                formType=AtomicFormTypeMeta(
                    type="INPUT_VARIABLE_PYTHON",
                ),
            ),
        ],
        outputList=[atomicMg.param("sentence_expanded_res", types="Str")],
    )
    def sentence_expand(
        model_select: LLMModelTypes = LLMModelTypes.DEEPSEEK_V3_2,
        prompt: str = PROMPT_SENTENCE_EXTEND,
        paragraph_text: str = "",
    ):
        """
        段落扩写

        :param paragraph_text: 段落

        :return:
            `str`, 扩写结果
        """

        # 生成提示词
        user_input = f"{prompt}\n\n段落：{paragraph_text}"
        model = model_select.value if isinstance(model_select, LLMModelTypes) else model_select

        # 向大模型发送请求（使用 chat_normal 替代 streamable_response）
        result = chat_normal(user_input=user_input, system_input="", model=model)
        s_content = result.replace("，", ",") if result else ""

        return s_content

    @staticmethod
    @atomicMg.atomic(
        "DocumentAI",
        inputList=[
            atomicMg.param(
                "prompt",
                formType=AtomicFormTypeMeta(
                    type="INPUT_PYTHON_TEXTAREAMODAL_VARIABLE",
                ),
            ),
            atomicMg.param(
                "paragraph_text",
                formType=AtomicFormTypeMeta(
                    type="INPUT_VARIABLE_PYTHON",
                ),
            ),
        ],
        outputList=[atomicMg.param("sentence_reduce_res", types="Str")],
    )
    def sentence_reduce(
        model_select: LLMModelTypes = LLMModelTypes.DEEPSEEK_V3_2,
        prompt: str = PROMPT_SENTENCE_REDUCE,
        paragraph_text: str = "",
    ):
        """
        段落缩写

        :param sentence_text: 段落

        :return:
            `str`, 缩写结果
        """

        # 生成提示词
        user_input = f"{prompt}\n\n段落：{paragraph_text}"
        model = model_select.value if isinstance(model_select, LLMModelTypes) else model_select

        # 向大模型发送请求（使用 chat_normal 替代 streamable_response）
        result = chat_normal(user_input=user_input, system_input="", model=model)
        s_content = result.replace("，", ",") if result else ""

        return s_content
